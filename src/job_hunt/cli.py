"""
AI智慧求职系统 - CLI入口
========================
纯CLI、AI驱动、本地运行的求职助手

用法:
  job-hunt init          初始化配置（AI引导式问答）
  job-hunt scan          扫描招聘平台抓取岗位
  job-hunt match         智能匹配评估岗位
  job-hunt eval <id>     深度评估某个岗位
  job-hunt resume <id>   生成定制化简历
  job-hunt apply <id>    自动投递
  job-hunt status        查看投递状态
  job-hunt chat          进入对话模式
  job-hunt verify <公司名> 公司信息交叉验证（五层）
  job-hunt parse <简历>   解析简历
"""

import asyncio
import json
import os
import sys
from pathlib import Path
from typing import Optional

import typer
from rich.prompt import Prompt, Confirm
from rich.console import Console

from . import __version__
from .db.database import Database
from .utils.config import Config
from .utils.display import (
    print_banner, print_info, print_success, print_warning, print_error, print_ai,
    print_status, display_job_table, display_job_detail, display_application_stats,
    display_application_table, display_resume_summary, console,
)
from .models.resume import Resume
from .models.job import Job
from .models.application import Application
try:
    from .ai.brain import AIBrain
    from .ai.verifier import verify_company_with_search, verify_company, VerifyResult
except ImportError:
    AIBrain = None
    verify_company = None
    verify_company_with_search = None
    VerifyResult = None

# ─── App 初始化 ───────────────────────────────────────────

app = typer.Typer(
    name="job-hunt",
    help="AI智慧求职系统 - 你只管说，AI来做",
    no_args_is_help=True,
)

# 全局变量（在命令中初始化）
_db: Optional[Database] = None
_config: Optional[Config] = None
_brain: Optional[AIBrain] = None


def get_db() -> Database:
    global _db
    if _db is None:
        _db = Database()
    return _db


def get_config() -> Config:
    global _config
    if _config is None:
        _config = Config()
    return _config


def get_brain() -> AIBrain:
    global _brain
    if AIBrain is None:
        print_error("litellm 未安装，AI功能不可用。请运行: pip install litellm")
        raise typer.Exit(1)
    if _brain is None:
        _brain = AIBrain(get_config())
    return _brain


def check_configured():
    """检查AI是否已配置"""
    if not get_config().is_configured:
        print_warning("⚠️  AI尚未配置！请先运行 job-hunt init 完成初始化。")
        raise typer.Exit(1)


# ─── HELP ────────────────────────────────────────────────

@app.callback(invoke_without_command=True)
def main(
    ctx: typer.Context,
    version: bool = typer.Option(False, "--version", "-v", help="显示版本号"),
):
    """AI智慧求职系统 - 纯CLI、本地运行、面向中国招聘市场"""
    if version:
        print_info(f"job-hunt v{__version__}")
        return
    if ctx.invoked_subcommand is None:
        print_banner()
        print_info("输入 [bold]job-hunt --help[/bold] 查看所有命令")
        print_info("输入 [bold]job-hunt init[/bold] 开始初始化")
        print()


# ─── INIT ────────────────────────────────────────────────

@app.command()
def init():
    """
    🚀 初始化配置 - AI引导式问答
    
    运行后将引导你完成：
    - 简历解析
    - AI API配置
    - 求职意向设置
    - 城市/薪资/岗位偏好
    """
    print_banner()
    print_info("欢迎使用 AI智慧求职系统！让我来帮你完成初始化...\n")

    config = get_config()
    db = get_db()

    # ── Step 1: AI API 配置 ──
    console.print("[bold cyan]📡 Step 1: AI 服务配置[/bold cyan]")
    console.print("支持 OpenAI / Claude / 通义千问 / 文心一言 / DeepSeek 等\n")

    provider = Prompt.ask(
        "  选择AI提供商",
        choices=["openai", "anthropic", "deepseek", "qwen", "ernie", "custom"],
        default=config.get("ai", "provider", "openai"),
    )
    config.set("ai", "provider", provider)

    # 根据provider推荐model
    default_models = {
        "openai": "gpt-4o-mini",
        "anthropic": "claude-3-5-haiku-20241022",
        "deepseek": "deepseek-chat",
        "qwen": "qwen-turbo",
        "ernie": "ernie-4.0-turbo",
    }
    default_model = default_models.get(provider, config.get("ai", "model", "gpt-4o-mini"))
    model = Prompt.ask("  模型名称", default=default_model)
    config.set("ai", "model", model)

    api_key = Prompt.ask(
        "  API Key（输入后不会显示）",
        password=True,
        default=config.get("ai", "api_key", ""),
    )
    config.set("ai", "api_key", api_key)

    api_base = Prompt.ask(
        "  API Base URL（OpenAI兼容可不填）",
        default=config.get("ai", "api_base", ""),
    )
    if api_base:
        config.set("ai", "api_base", api_base)

    print_success("AI 配置完成\n")

    # ── Step 2: 简历解析 ──
    console.print("[bold cyan]📄 Step 2: 简历解析[/bold cyan]")

    brain = AIBrain(config)
    resume = db.get_resume()

    has_existing = resume is not None and resume.name
    if has_existing:
        console.print(f"  已找到简历: {resume.name} | {resume.education_level} - {resume.major}")
        if not Confirm.ask("  是否重新解析简历？", default=False):
            print_info("跳过简历解析\n")
        else:
            has_existing = False

    if not has_existing:
        resume_path = Prompt.ask("  简历文件路径（PDF/DOCX/TXT）", default="")
        if resume_path and os.path.exists(resume_path):
            print_info("正在解析简历...")
            try:
                # 读取文本
                ext = Path(resume_path).suffix.lower()
                if ext == ".pdf":
                    try:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(resume_path)
                        raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                    except ImportError:
                        print_warning("PyPDF2 未安装，尝试作为纯文本读取")
                        with open(resume_path, "r", encoding="utf-8") as f:
                            raw_text = f.read()
                elif ext in (".docx", ".doc"):
                    try:
                        from docx import Document
                        doc = Document(resume_path)
                        raw_text = "\n".join(p.text for p in doc.paragraphs)
                    except ImportError:
                        print_warning("python-docx 未安装，尝试作为纯文本读取")
                        with open(resume_path, "rb") as f:
                            raw_text = f.read().decode("utf-8", errors="ignore")
                else:
                    with open(resume_path, "r", encoding="utf-8") as f:
                        raw_text = f.read()

                # LLM解析
                parsed = brain.parse_resume(raw_text)
                resume = Resume(
                    name=parsed.get("name", ""),
                    phone=parsed.get("phone", ""),
                    email=parsed.get("email", ""),
                    wechat=parsed.get("wechat", ""),
                    education_level=parsed.get("education_level", ""),
                    university=parsed.get("university", ""),
                    major=parsed.get("major", ""),
                    graduation_year=parsed.get("graduation_year", 0),
                    projects=json.dumps(parsed.get("projects", []), ensure_ascii=False),
                    skills=parsed.get("skills", ""),
                    work_years=parsed.get("work_years", 0),
                    desired_city=parsed.get("desired_city", ""),
                    desired_position=parsed.get("desired_position", ""),
                    salary_min=parsed.get("salary_min", 0),
                    salary_max=parsed.get("salary_max", 0),
                    raw_text=raw_text[:5000],
                    raw_file_path=resume_path,
                )
                db.save_resume(resume)
                print_success(f"简历解析完成: {resume.summary()}")
                display_resume_summary(resume)
            except Exception as e:
                print_warning(f"简历解析出错: {e}")
                print_info("可以稍后手动设置")
        else:
            print_info("跳过简历解析（可稍后通过 job-hunt init 重新设置）")
            # 手动输入基本信息
            name = Prompt.ask("  姓名", default="")
            if name:
                resume = Resume(name=name)
                resume.education_level = Prompt.ask("  学历（本科/硕士/博士）", default="硕士")
                resume.major = Prompt.ask("  专业", default="")
                resume.university = Prompt.ask("  学校", default="")
                resume.skills = Prompt.ask("  技能（逗号分隔）", default="")
                db.save_resume(resume)

    print()

    # ── Step 3: 求职意向 ──
    console.print("[bold cyan]🎯 Step 3: 求职意向[/bold cyan]")

    cities = Prompt.ask("  意向城市（逗号分隔）", default=config.get("preferences", "cities", "广西"))
    config.set("preferences", "cities", cities)

    position = Prompt.ask("  意向岗位", default=config.get("preferences", "position", ""))
    config.set("preferences", "position", position)

    industry = Prompt.ask("  意向行业", default=config.get("preferences", "industry", ""))
    config.set("preferences", "industry", industry)

    salary_min = Prompt.ask("  最低月薪（元）", default=str(config.get("preferences", "salary_min", 8000)))
    config.set("preferences", "salary_min", int(salary_min))

    salary_max = Prompt.ask("  最高月薪（元）", default=str(config.get("preferences", "salary_max", 15000)))
    config.set("preferences", "salary_max", int(salary_max))

    print()

    # ── Step 4: 关键词 ──
    console.print("[bold cyan]🔑 Step 4: 搜索关键词[/bold cyan]")
    keywords = Prompt.ask(
        "  搜索关键词（逗号分隔，精准描述你的专业方向）",
        default=config.get("scanner", "keywords", ""),
    )
    config.set("scanner", "keywords", keywords)

    print()
    console.print(
        "[bold green]🎉 初始化完成！[/bold green]\n"
        f"\n  求职画像: {resume.education_level} | {resume.major} | {position} | {cities}"
        f"\n  关键词: {keywords}\n"
        f"\n  下一步: [bold]job-hunt scan[/bold] 开始扫描岗位\n"
    )


# ─── SCAN ────────────────────────────────────────────────

@app.command()
def scan(
    city: Optional[str] = typer.Option(None, "--city", "-c", help="城市搜索（如：广西、南宁）"),
    keyword: Optional[str] = typer.Option(None, "--keyword", "-k", help="搜索关键词"),
    platform: str = typer.Option("all", "--platform", "-p", help="招聘平台：all/gxrc/job51/boss/bing"),
    max_pages: int = typer.Option(3, "--pages", "-n", help="每站最大页数"),
    headless: bool = typer.Option(True, "--headless/--no-headless", help="BOSS直聘模式：是否无头浏览器"),
    debug: bool = typer.Option(False, "--debug", help="保存原始HTML用于调试选择器"),
):
    """
    🔍 真实岗位扫描 - 直连招聘网站抓取，非LLM编造

    渠道覆盖：广西人才网(gxrc) | 前程无忧(51job) | BOSS直聘 | Bing补充

    示例:
      job-hunt scan -k "环境信息系统 数据分析" -c 广西
      job-hunt scan -k "环保 环境工程" --platform gxrc
      job-hunt scan -k "Python" --platform boss --headless
      job-hunt scan -k "数据分析" -c 南宁 --debug
    """
    check_configured()
    config = get_config()
    db = get_db()
    resume = db.get_resume()

    city = city or config.cities or ""
    keyword = keyword or config.keywords
    if not keyword:
        print_error("请指定搜索关键词 (--keyword)")
        raise typer.Exit(1)

    print_banner()

    from .scrapers.job51 import Job51Scraper
    from .scrapers.bing import bing_job_search

    all_new_jobs: list = []

    # 确定要运行的抓取器
    if platform == "all":
        targets = ["gxrc", "job51", "bing"]
    else:
        if platform not in ("gxrc", "job51", "boss", "bing"):
            print_error(f"未知平台: {platform}，可选: all/gxrc/job51/boss/bing")
            raise typer.Exit(1)
        targets = [platform]

    print_status(f"🔍 真实岗位扫描 | 渠道: {', '.join(targets)} | 城市: {city or '全国'} | 关键词: {keyword}")
    print()

    for plat in targets:
        if plat == "gxrc":
            # ─── GXRC: Playwright 浏览器模式（SPA网站） ───
            scraper_name = "广西人才网"
            print_status(f"  🔍 {scraper_name} (浏览器模式)...")
            from .scrapers.gxrc import GxrcScraper

            async def _scan_gxrc():
                scraper = GxrcScraper(headless=headless, debug=debug)
                try:
                    return await scraper.search(keyword=keyword, city=city, max_pages=max_pages)
                finally:
                    await scraper.close()
            try:
                jobs = asyncio.run(_scan_gxrc())
            except ImportError as e:
                print_warning(f"  {scraper_name}: {e}")
                continue
            except Exception as e:
                print_warning(f"  {scraper_name}: 抓取出错 - {e}")
                continue

        elif plat == "job51":
            scraper_name = "前程无忧"
            print_status(f"  🔍 {scraper_name}...")
            try:
                scraper = Job51Scraper(debug=debug)
                try:
                    jobs = scraper.search(keyword=keyword, city=city, max_pages=max_pages)
                finally:
                    scraper.close()
            except Exception as e:
                print_warning(f"  {scraper_name}: 抓取出错 - {e}")
                continue

        elif plat == "boss":
            scraper_name = "BOSS直聘"
            print_status(f"  🔍 {scraper_name} (浏览器模式)...")
            from .scrapers.boss import BossScraper

            async def _scan_boss():
                scraper = BossScraper(headless=headless)
                try:
                    return await scraper.search(keyword=keyword, city=city, max_pages=max_pages)
                finally:
                    await scraper.close()
            try:
                jobs = asyncio.run(_scan_boss())
            except ImportError as e:
                print_warning(f"  {scraper_name}: {e}")
                continue
            except Exception as e:
                print_warning(f"  {scraper_name}: 抓取出错 - {e}")
                continue

        elif plat == "bing":
            scraper_name = "Bing搜索引擎"
            print_status(f"  🔍 {scraper_name}...")
            jobs = bing_job_search(keyword=keyword, city=city, max_results=max_pages * 5)

        # 保存到数据库
        saved = 0
        for job in jobs:
            if job.title:
                db.save_job(job)
                saved += 1
                all_new_jobs.append(job)

        if saved > 0:
            print_success(f"  ✅ {scraper_name}: {saved} 条岗位")
        else:
            print_status(f"  ⚪ {scraper_name}: 无结果（加 --debug 保存HTML调试选择器）")

    # 统计
    all_count = db.get_job_count()
    city_count = db.get_job_count(city=city) if city else all_count

    print()
    print_success(f"✅ 扫描完成 | 本次新增: {len(all_new_jobs)} 个岗位 | "
                  f"城市相关: {city_count} 个 | 总岗位库: {all_count} 个")

    if all_new_jobs:
        display_job_table(all_new_jobs[:20], title="📋 本次新增岗位")

    print_info("下一步: [bold]job-hunt match[/bold] 开始智能匹配")


# ─── MATCH ───────────────────────────────────────────────

@app.command()
def match(
    limit: Optional[int] = typer.Option(None, "--limit", "-n", help="匹配岗位数量"),
    min_score: float = typer.Option(0, "--min-score", "-m", help="最低匹配度（0-100）"),
    city: Optional[str] = typer.Option(None, "--city", "-c", help="城市筛选"),
):
    """
    🎯 智能匹配 - AI分析岗位与简历的匹配度
    
    示例:
      job-hunt match              # 匹配所有未评估的岗位
      job-hunt match -n 20 -m 60  # 匹配20个，只看60%以上
    """
    check_configured()
    db = get_db()
    brain = get_brain()
    config = get_config()

    resume = db.get_resume()
    if not resume:
        print_error("请先运行 job-hunt init 设置简历")
        raise typer.Exit(1)

    city = city or config.cities or ""
    jobs = db.get_jobs(limit=limit or 200, city=city, active_only=True)
    # 只评估未匹配过的
    unrated = [j for j in jobs if j.match_score == 0]

    if not unrated:
        # 没有未匹配的，显示已匹配的
        rated = [j for j in jobs if j.match_score >= min_score]
        rated.sort(key=lambda j: j.match_score, reverse=True)
        print_banner()
        print_info(f"所有岗位已完成匹配，共 {len(rated)} 个匹配结果")
        display_job_table(rated[:30], title="📊 最佳匹配")
        return

    print_banner()
    print_info(f"🤖 开始AI匹配分析 | 待评估: {len(unrated)} 个岗位 | 城市: {city or '全国'}")

    resume_dict = resume.to_dict()
    matched = []

    for i, job in enumerate(unrated, 1):
        print_status(f"[{i}/{len(unrated)}] 分析: {job.title} - {job.company}")

        try:
            result = brain.match_job(resume_dict, job.to_dict())
            score = result.get("match_score", 0)
            reasons = result.get("reasons", [])
            gaps = result.get("gaps", [])
            suggestion = result.get("suggestions", "")

            reason_text = f"✅ 匹配点: {'; '.join(reasons[:3])}" if reasons else ""
            gap_text = f"⚠️ 差距: {'; '.join(gaps[:2])}" if gaps else ""

            db.update_job_match(
                job.id or 0,
                match_score=score,
                match_detail=json.dumps(result, ensure_ascii=False),
            )

            # 如果匹配度较高，生成推荐理由
            if score >= 70:
                recommend = suggestion or f"匹配度{score:.0f}%，技能和经验与岗位要求契合"
            elif score >= 50:
                recommend = suggestion or f"匹配度{score:.0f}%，部分匹配，建议进一步了解"
            else:
                recommend = suggestion or ""

            # 更新推荐理由
            job.match_score = score
            job.recommend_reason = recommend
            matched.append(job)

        except Exception as e:
            print_warning(f"匹配出错: {e}")
            continue

    # 排序显示
    matched.sort(key=lambda j: j.match_score, reverse=True)

    print_success(f"\n✅ 匹配完成！共评估 {len(matched)} 个岗位")

    if matched:
        top = matched[:min(15, len(matched))]
        display_job_table(top, title="🏆 匹配度排名 TOP")

        print_info("\n💡 高匹配度岗位建议优先投递（匹配度 ≥ 70%）")
        high_match = [j for j in matched if j.match_score >= 70]
        if high_match:
            for j in high_match[:5]:
                print_info(f"  #{j.id}: {j.title} @ {j.company} [{j.match_score:.0f}%]")
        else:
            print_info("  暂无高匹配度岗位，建议扩大搜索范围或调整关键词")

    print_info("\n下一步: [bold]job-hunt eval <id>[/bold] 深度评估某岗位")


# ─── EVAL ────────────────────────────────────────────────

@app.command()
def eval(
    job_id: int = typer.Argument(..., help="岗位ID（从 match 结果中获取）"),
):
    """
    📊 深度评估 - AI六维分析某个岗位
    
    评估维度：岗位匹配度 | 职级定位 | 薪资水平 | 公司质量 | 成长空间 | 投递优先级
    
    示例:
      job-hunt eval 1       # 评估ID为1的岗位
    """
    check_configured()
    db = get_db()
    brain = get_brain()

    job = db.get_job_by_id(job_id)
    if not job:
        print_error(f"未找到岗位 #{job_id}")
        raise typer.Exit(1)

    resume = db.get_resume()
    if not resume:
        print_error("请先运行 job-hunt init 设置简历")
        raise typer.Exit(1)

    print_banner()
    print_info(f"🔬 深度评估: {job.title} - {job.company}")

    result = brain.evaluate_job(job.to_dict(), resume.to_dict())

    # 显示评估结果
    from rich.panel import Panel
    from rich.table import Table

    overall = result.get("overall_score", "N/A")
    priority = result.get("投递优先级", "N/A")

    console.print()
    console.print(Panel.fit(
        f"[bold white]{job.title}[/bold white]\n[dim]{job.company} | {job.city} | {job.salary_range_display}[/dim]",
        border_style="cyan",
        title=f"综合评分: [bold yellow]{overall}[/bold yellow]",
    ))

    # 评估维度表
    eval_table = Table(box=None, padding=(0, 2))
    eval_table.add_column("维度", style="dim", width=14)
    eval_table.add_column("评分", style="bold", width=8)
    eval_table.add_column("说明", style="white")

    dimensions = ["岗位匹配度", "职级定位", "薪资水平", "公司质量", "成长空间"]
    for dim in dimensions:
        dim_data = result.get(dim, {})
        score = dim_data.get("score", "-")
        comment = dim_data.get("comment", "")
        color = "green" if score.startswith("A") else "yellow" if score.startswith("B") else "red"
        eval_table.add_row(dim, f"[{color}]{score}[/{color}]", comment)

    eval_table.add_row("投递优先级", f"[bold magenta]{priority}[/bold magenta]", "")

    console.print(eval_table)

    # 面试准备
    questions = result.get("interview_questions", [])
    if questions:
        console.print("\n[bold yellow]🎤 面试可能问到:[/bold yellow]")
        for q in questions:
            console.print(f"  • {q}")

    advice = result.get("advice", "")
    if advice:
        console.print(f"\n[bold cyan]💡 准备建议:[/bold cyan] {advice}")

    # 保存评估结果
    db.update_job_eval(job_id, overall, json.dumps(result, ensure_ascii=False))

    print_info(f"\n下一步: [bold]job-hunt resume {job_id}[/bold] 生成定制简历")


# ─── RESUME ──────────────────────────────────────────────

@app.command()
def resume(
    job_id: int = typer.Argument(..., help="岗位ID"),
    output_format: str = typer.Option("pdf", "--format", "-f", help="输出格式: pdf/md/html"),
):
    """
    📄 生成定制化简历 - AI根据JD优化简历
    
    示例:
      job-hunt resume 1           # 生成PDF简历
      job-hunt resume 1 -f md     # 生成Markdown简历
    """
    check_configured()
    db = get_db()
    brain = get_brain()

    job = db.get_job_by_id(job_id)
    if not job:
        print_error(f"未找到岗位 #{job_id}")
        raise typer.Exit(1)

    resume = db.get_resume()
    if not resume:
        print_error("请先运行 job-hunt init 设置简历")
        raise typer.Exit(1)

    print_banner()
    print_info(f"📝 生成定制简历 | 目标: {job.title} - {job.company}")

    # LLM生成优化后的简历内容
    print_info("AI正在分析岗位JD并优化简历...")
    md_content = brain.generate_resume(resume.to_dict(), job.to_dict())

    # 保存
    safe_title = "".join(c for c in job.title if c.isalnum() or c in " _-")[:30]
    safe_company = "".join(c for c in job.company if c.isalnum() or c in " _-")[:20]
    base_name = f"{safe_company}_{safe_title}_简历"
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    md_path = output_dir / f"{base_name}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print_success(f"简历已保存: {md_path}")

    # 生成PDF
    if output_format == "pdf":
        try:
            from weasyprint import HTML
            html_content = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 0 auto; padding: 40px; color: #333; }}
h1 {{ color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 8px; }}
h2 {{ color: #1a73e8; margin-top: 24px; }}
h3 {{ margin-top: 16px; }}
ul {{ padding-left: 20px; }}
li {{ margin: 4px 0; }}
</style></head><body>
{_md_to_html(md_content)}
</body></html>"""
            pdf_path = output_dir / f"{base_name}.pdf"
            HTML(string=html_content).write_pdf(str(pdf_path))
            print_success(f"PDF简历已生成: {pdf_path}")
        except ImportError:
            print_warning("weasyprint未安装，跳过PDF生成。请安装: pip install weasyprint")
        except Exception as e:
            print_warning(f"PDF生成失败: {e}")

    print_info(f"\n简历文件在 output/ 目录下，可直接使用")


def _md_to_html(md_text: str) -> str:
    """简易Markdown转HTML（避免额外依赖）"""
    import re
    lines = md_text.split("\n")
    html_lines = []
    in_list = False

    for line in lines:
        if line.startswith("### "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h3>{line[4:]}</h3>")
        elif line.startswith("## "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith("# "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{line[2:]}</li>")
        elif line.strip() == "":
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("")
        else:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            # Bold
            line = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', line)
            html_lines.append(f"<p>{line}</p>")

    if in_list:
        html_lines.append("</ul>")
    return "\n".join(html_lines)


# ─── APPLY ───────────────────────────────────────────────

@app.command()
def apply(
    job_id: int = typer.Argument(..., help="岗位ID"),
):
    """
    🚀 自动投递 - AI帮你投递简历
    
    需要先在浏览器中登录招聘平台
    
    示例:
      job-hunt apply 1       # 自动投递ID为1的岗位
    """
    check_configured()
    db = get_db()
    brain = get_brain()
    config = get_config()

    job = db.get_job_by_id(job_id)
    if not job:
        print_error(f"未找到岗位 #{job_id}")
        raise typer.Exit(1)

    resume = db.get_resume()
    if not resume:
        print_error("请先运行 job-hunt init 设置简历")
        raise typer.Exit(1)

    print_banner()
    print_info(f"🚀 准备投递: {job.title} - {job.company} ({job.platform_display})")

    # 生成打招呼语
    greeting = brain.generate_greeting(resume.to_dict(), job.to_dict())
    print_info(f"💬 打招呼语: {greeting}")

    # 确认
    from rich.prompt import Confirm
    if not Confirm.ask("确认投递？", default=True):
        print_info("已取消投递")
        return

    # 执行投递
    if job.platform == "boss":
        from .applier.auto_apply import AutoApplier

        async def _apply():
            applier = AutoApplier()
            try:
                success = await applier.apply_boss(job, "", greeting)
                return success
            finally:
                await applier.close()

        success = asyncio.run(_apply())
    else:
        print_warning(f"平台 {job.platform} 暂不支持自动投递")
        print_info(f"请手动投递: {job.source_url}")
        success = False

    # 记录投递
    app_record = Application(
        job_id=job.id or 0,
        job_title=job.title,
        company=job.company,
        platform=job.platform,
        status="applied",
        greeting=greeting,
        notes="自动投递" if success else "需手动投递",
    )
    db.save_application(app_record)

    if success:
        print_success(f"✅ 已投递: {job.title} - {job.company}")
        print_info("可以在 job-hunt status 中查看投递状态")
    else:
        print_info("浏览器窗口将继续保持打开，请手动完成操作")


# ─── STATUS ──────────────────────────────────────────────

@app.command()
def status():
    """
    📊 查看投递状态和历史记录
    
    示例:
      job-hunt status
    """
    db = get_db()
    print_banner()

    stats = db.get_application_stats()
    display_application_stats(stats)

    apps = db.get_applications(limit=30)
    display_application_table(apps)

    if not apps:
        print_info("还没有投递记录，试试 [bold]job-hunt match[/bold] 找到合适的岗位！")

    # 岗位库统计
    job_count = db.get_job_count()
    resume = db.get_resume()
    print()
    console.print(
        f"[dim]岗位库: {job_count} 个 | 简历: {'已导入' if resume else '未导入'}[/dim]\n"
    )


# ─── CHAT ────────────────────────────────────────────────

@app.command()
def chat():
    """
    💬 对话模式 - 自由向AI提问求职相关问题
    
    示例:
      job-hunt chat
      
      进入后在 >>> 后输入问题即可
      输入 quit / exit / 退出 结束对话
    """
    check_configured()
    db = get_db()
    brain = get_brain()
    config = get_config()

    resume = db.get_resume()
    resume_dict = resume.to_dict() if resume else {}
    job_count = db.get_job_count()

    print_banner()
    console.print("[bold cyan]💬 AI求职助手对话模式[/bold cyan]")
    console.print("你可以问我任何求职相关的问题，输入 [bold]quit[/bold] / [bold]exit[/bold] / [bold]退出[/bold] 结束对话\n")

    context = {
        "resume": resume_dict,
        "job_count": job_count,
        "cities": config.cities,
        "keywords": config.keywords,
    }

    console.print(f"[dim]求职画像: {config.cities} | {config.keywords} | 岗位库: {job_count}个[/dim]\n")

    history = []
    while True:
        try:
            user_input = Prompt.ask("[bold cyan]>>>[/bold cyan]")
        except (KeyboardInterrupt, EOFError):
            break

        if user_input.strip().lower() in ("quit", "exit", "退出", "q"):
            print_info("再见！祝你求职顺利 🎉")
            break

        if not user_input.strip():
            continue

        history.append({"role": "user", "content": user_input})

        try:
            # 构建增强上下文
            full_context = {
                **context,
                "history": history[-5:],  # 最近5轮对话
            }
            response = brain.chat(user_input, full_context)
            console.print(f"\n[green]🤖[/green] {response}\n")
            history.append({"role": "assistant", "content": response})
        except Exception as e:
            print_warning(f"AI响应出错: {e}")


# ─── PARSE ───────────────────────────────────────────────

@app.command()
def parse(
    file_path: str = typer.Argument(..., help="简历文件路径（PDF/DOCX/TXT）"),
):
    """
    📄 解析简历 - 单独运行简历解析
    
    示例:
      job-hunt parse resume.pdf
    """
    _ = get_config()
    db = get_db()
    brain = get_brain()

    if not os.path.exists(file_path):
        print_error(f"文件不存在: {file_path}")
        raise typer.Exit(1)

    print_banner()
    print_info(f"解析简历: {file_path}")

    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                print_warning("PyPDF2未安装，使用纯文本读取")
                with open(file_path, "r", encoding="utf-8") as f:
                    raw_text = f.read()
        elif ext in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document(file_path)
                raw_text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                with open(file_path, "rb") as f:
                    raw_text = f.read().decode("utf-8", errors="ignore")
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()

        parsed = brain.parse_resume(raw_text)
        resume = Resume(
            name=parsed.get("name", ""),
            phone=parsed.get("phone", ""),
            email=parsed.get("email", ""),
            wechat=parsed.get("wechat", ""),
            education_level=parsed.get("education_level", ""),
            university=parsed.get("university", ""),
            major=parsed.get("major", ""),
            graduation_year=parsed.get("graduation_year", 0),
            projects=json.dumps(parsed.get("projects", []), ensure_ascii=False),
            skills=parsed.get("skills", ""),
            work_years=parsed.get("work_years", 0),
            desired_city=parsed.get("desired_city", ""),
            desired_position=parsed.get("desired_position", ""),
            salary_min=parsed.get("salary_min", 0),
            salary_max=parsed.get("salary_max", 0),
            raw_text=raw_text[:5000],
            raw_file_path=file_path,
        )
        db.save_resume(resume)
        display_resume_summary(resume)
        print_success("简历已存入本地数据库")

    except Exception as e:
        print_error(f"解析失败: {e}")


# ─── VERIFY ──────────────────────────────────────────────

@app.command()
def verify(
    company_name: str = typer.Argument(..., help="公司全称（如：广西环保产业投资集团有限公司）"),
    direction: Optional[str] = typer.Option(None, "--direction", "-d", help="关注方向（如：环境信息系统）"),
    deep: bool = typer.Option(False, "--deep", help="深度模式：会进行多次网络搜索交叉验证"),
):
    """
    🔍 公司信息交叉验证 - 五层验证确保信息真实

    验证层次：
    ① 工商注册（是否合法注册） ② 官方渠道（官网/主管单位）
    ③ 招聘真实性（官方发布 vs 第三方） ④ 业务/部门真实性
    ⑤ 风险扫描（失信/被执行/行政处罚）

    示例：
      job-hunt verify "广西环保产业投资集团有限公司"
      job-hunt verify "广西环科院" -d "环境信息系统"
      job-hunt verify "某公司" --deep  # 深度验证，多源交叉搜索
    """
    check_configured()
    config = get_config()
    direction = direction or config.get("preferences", "position", "环境信息系统")

    print_banner()
    print_info(f"🔍 开始交叉验证: [bold]{company_name}[/bold]")
    print_info(f"关注方向: {direction}")
    print()

    with console.status("[cyan]正在进行五层交叉验证...[/cyan]", spinner="dots"):
        if not verify_company_with_search:
            print_error("验证模块不可用：litellm 未安装")
            raise typer.Exit(1)

        if deep:
            # 深度模式：多源搜索
            search_text = _deep_search_company(company_name)
            result = verify_company_with_search(company_name, search_text, direction)
        else:
            # 快速模式：LLM 知识库判断
            result = verify_company(company_name, direction)

    # ── 显示验证结果 ──
    _display_verify_result(result, config)

    # ── 把验证的公司信息写入本地库（供 match/eval 使用）──
    db = get_db()
    from .models.job import Job
    job = Job(
        title=f"[待扫描] {direction}相关岗位",
        company=company_name,
        city=config.cities,
        platform="verified",
        source_url=result.website_url,
        education="本科及以上",
        recommend_reason=f"验证结果: {result.verdict} | 得分: {result.overall_score}",
        eval_score=str(result.overall_score),
        eval_detail=json.dumps({
            "verdict": result.verdict,
            "risk_level": result.risk_level,
            "is_state_owned": result.is_state_owned,
            "registered": result.registered,
            "business_lines": result.business_lines,
        }, ensure_ascii=False),
    )
    saved_id = db.save_job(job)
    print_info(f"已存入本地数据库 (ID: {saved_id})")


def _deep_search_company(company_name: str) -> str:
    """深度搜索：多源交叉搜索公司信息（真正调用搜索引擎）"""
    try:
        from .scrapers.engine import verify_search
        print_status(f"  🔍 多源搜索中: 天眼查 → 官网 → 招聘 → 风险...")
        result = verify_search(company_name)
        if result.strip():
            return result
        else:
            print_warning("  ⚠️ 搜索无结果，切换到LLM知识库模式")
            return f'未搜索到"{company_name}"的公开信息，请基于知识库判断。'
    except ImportError as e:
        print_warning(f"  ⚠️ 搜索模块不可用: {e}")
        return f'无法搜索，请基于知识库判断"{company_name}"。'


def _display_verify_result(result: VerifyResult, config):
    """显示验证结果面板"""
    from rich.panel import Panel

    # 标题
    verdict_color = {"可信": "green", "待确认": "yellow", "存疑": "red", "排除": "red"}
    color = verdict_color.get(result.verdict, "yellow")

    console.print()
    console.print(Panel.fit(
        f"[bold white]{result.company_name}[/bold white]\n"
        f"综合评分: [{color}]{result.overall_score}/100[/{color}]   |   "
        f"判定: [{color}]{result.verdict_display}[/{color}]",
        border_style=color,
        title="🔍 交叉验证报告",
    ))

    # 五层验证详情
    from rich.table import Table
    table = Table(box=None, padding=(0, 2), show_header=False)
    table.add_column("层", style="dim", width=3)
    table.add_column("项", width=16)
    table.add_column("结果", style="white")

    # 第一层
    reg_status = "✅ 已注册" if result.registered else "❌ 未确认"
    table.add_row("①", "工商注册", f"{reg_status} | {result.registration_info}" if result.registration_info else reg_status)
    table.add_row("", "企业性质", result.is_state_owned or "未知")

    # 第二层
    web_status = f"✅ {result.website_url}" if result.has_website else "⚠️ 未确认"
    table.add_row("②", "官网", web_status)

    # 第三层
    hire_status = "✅ 确认" if result.hiring_confirmed else ("⚠️ 待确认" if result.has_recent_hiring else "❌ 无信息")
    table.add_row("③", "近期招聘", f"{hire_status} | {result.hiring_sources}" if result.hiring_sources else hire_status)

    # 第四层
    dept_status = "✅ 有相关部门" if result.has_target_dept else "⚠️ 未确认"
    table.add_row("④", "业务方向", result.business_lines or "未确认")
    table.add_row("", "对口部门", dept_status)

    # 第五层
    risk_color_map = {"绿色": "green", "黄色": "yellow", "红色": "red"}
    rc = risk_color_map.get(result.risk_level, "yellow")
    table.add_row("⑤", "风险等级", f"[{rc}]{'🟢' if result.risk_level=='绿色' else '🟡' if result.risk_level=='黄色' else '🔴'} {result.risk_level}[/{rc}]")
    if result.risk_details and result.risk_details != "未发现":
        table.add_row("", "风险说明", result.risk_details)

    console.print(table)

    # 关键证据 & 警告
    if result.evidence:
        console.print(f"\n📋 [bold]证据:[/bold] {result.evidence}")

    if result.warnings:
        console.print("\n⚠️ [bold yellow]警告:[/bold yellow]")
        for w in result.warnings:
            console.print(f"  • {w}")

    # 建议操作
    console.print("\n💡 [bold]建议操作:[/bold]")
    if result.verdict == "可信":
        console.print("  ✅ 可放心投递，建议访问官网确认最新招聘信息")
        console.print(f"  📎 广西人才网: https://www.gxrc.com 搜索 \"{result.company_name}\"")
    elif result.verdict == "待确认":
        console.print("  ⚠️ 建议手动查天眼查确认工商信息: https://www.tianyancha.com")
        console.print("  ⚠️ 建议访问企业官网确认是否有招聘")
    elif result.verdict == "存疑":
        console.print("  🔴 强烈建议手动核实后再决定是否投递")
        console.print("  🔴 查天眼查 + 国家企业信用信息公示系统: http://www.gsxt.gov.cn")
    elif result.verdict == "排除":
        console.print("  ⚫ 不建议投递，存在严重风险信号")


# ─── ENTRY ───────────────────────────────────────────────

def main_cli():
    """程序入口"""
    app()


if __name__ == "__main__":
    main_cli()
